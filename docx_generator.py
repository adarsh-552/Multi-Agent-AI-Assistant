from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION


def create_resume_docx(resume_text, output_path):

    document = Document()

    # ==================================================
    # PAGE SETTINGS
    # ==================================================

    section = document.sections[0]

    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)


    # ==================================================
    # DEFAULT FONT
    # ==================================================

    styles = document.styles

    normal_style = styles["Normal"]

    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9.5)


    # ==================================================
    # SECTION HEADINGS
    # ==================================================

    headings = {
        "PROFESSIONAL SUMMARY",
        "SUMMARY",
        "TECHNICAL SKILLS",
        "SKILLS",
        "PROFESSIONAL SKILLS",
        "TOOLS",
        "PROJECTS",
        "EXPERIENCE",
        "EDUCATION",
        "CERTIFICATIONS",
        "ACHIEVEMENTS",
        "ACHIEVEMENTS & EXTRACURRICULARS",
        "EXTRACURRICULAR ACTIVITIES",
    }


    # ==================================================
    # CLEAN TEXT
    # ==================================================

    lines = resume_text.splitlines()

    cleaned_lines = []

    for line in lines:

        line = line.strip()

        if line:

            cleaned_lines.append(line)


    if not cleaned_lines:

        document.add_paragraph(
            "No resume content available."
        )

        document.save(output_path)

        return output_path


    # ==================================================
    # CANDIDATE NAME
    # ==================================================

    name_paragraph = document.add_paragraph()

    name_paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    name_paragraph.paragraph_format.space_after = Pt(2)

    name_run = name_paragraph.add_run(
        cleaned_lines[0]
    )

    name_run.bold = True
    name_run.font.name = "Arial"
    name_run.font.size = Pt(18)


    # ==================================================
    # PROCESS REMAINING LINES
    # ==================================================

    for line in cleaned_lines[1:]:

        upper_line = line.upper()


        # ==================================================
        # SECTION HEADING
        # ==================================================

        if upper_line in headings:

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_before = Pt(7)
            paragraph.paragraph_format.space_after = Pt(2)

            run = paragraph.add_run(
                upper_line
            )

            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)


            # Add underline
            run.underline = True

            continue


        # ==================================================
        # BULLET POINT
        # ==================================================

        if line.startswith(
            ("•", "●", "-", "▪", "○")
        ):

            clean_line = line.lstrip(
                "•●-▪○ "
            )

            paragraph = document.add_paragraph(
                style="List Bullet"
            )

            paragraph.paragraph_format.left_indent = (
                Inches(0.18)
            )

            paragraph.paragraph_format.first_line_indent = (
                Inches(-0.10)
            )

            paragraph.paragraph_format.space_after = Pt(1)

            run = paragraph.add_run(
                clean_line
            )

            run.font.name = "Arial"
            run.font.size = Pt(9.2)

            continue


        # ==================================================
        # NORMAL TEXT
        # ==================================================

        paragraph = document.add_paragraph()

        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(
            line
        )

        run.font.name = "Arial"
        run.font.size = Pt(9.2)


    # ==================================================
    # SAVE DOCX
    # ==================================================

    document.save(
        output_path
    )

    return output_path